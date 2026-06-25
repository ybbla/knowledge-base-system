"""DOCX 文档解析器
使用 python-docx .docx 文件解析为统一ParsedElement Asset支持标题、段落、列表、表格、内联图片、超链接和嵌入资源的提取"""

import hashlib
import io
import logging
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path, PurePosixPath

from docx import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from app.core.paths import resolve_file_uri
from app.core.models import (
    Asset,
    AssetData,
    AssetStatus,
    AssetType,
    Document,
    ElementType,
    ParsedElement,
    SourceLocation,
    compute_hash,
)
from parsers.base import DocumentParser, ParseResult, _BaseParseState
from parsers.utils import (
    ATTACHMENT_EXTENSIONS,
    classify_link,
    classify_link_text,
)

logger = logging.getLogger(__name__)

# 已知图片扩展名（用于链接 URL 分类，与 MarkdownParser 保持一致）
_IMAGE_EXTENSIONS: set[str] = {
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg", ".tiff", ".tif",
}


class DocxParser(DocumentParser):
    """将 DOCX 文档解析为 ParsedElement 和 Asset。

    支持段落、标题、列表、表格、内联图片、超链接和嵌入资源。
    """

    SUPPORTED_TYPES = {"docx"}

    # 多语言标题样式关键词集合。
    HEADING_KEYWORDS = {
        "heading", "head",          # 英文
        "标题",                      # 中文
        "titre", "title",           # 法文 / 英文变体
        "uberschrift", "überschrift",  # 德文
        "titulo", "título",         # 西班牙文 / 葡萄牙文
        "intestazione",             # 意大利文
        "見出し",                    # 日文
    }

    def supports(self, source_type: str) -> bool:
        return source_type.lower() in self.SUPPORTED_TYPES

    # ── 主解析入──────────────────────────────────────────────

    def parse(self, doc: Document, content: bytes | str) -> ParseResult:
        """主解析入口：解析 DOCX 文档的结构化内容
        新流程：预提取图遍历 body flush 提取视频 回填关联 计算哈希        """
        if isinstance(content, str):
            content = content.encode("utf-8")
        docx = DocxDocument(io.BytesIO(content))
        state = _DocxParseState(doc.doc_id, doc.version)

        # 1. 预提取所有嵌入图片到 asset map
        self._build_image_asset_map(content, doc, state)

        # 2. 按文档顺序遍历所有正文元素，不处理页眉和页脚。
        for child in docx.element.body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                self._process_paragraph(child, docx, state)
            elif tag == "tbl":
                self._process_table(child, docx, state)

        # 3. 完成解析并获取全部元素。
        elements = state.flush_elements()

        # 4. 回填 Asset.element_id 关联
        self._link_assets_to_elements(elements, state.assets)

        # 6. 计算文档内容哈希
        doc.source_hash = compute_hash(content)

        return ParseResult(doc=doc, elements=elements, assets=state.assets)

    # ── 标题检────────────────────────────────────────────────

    @classmethod
    def _detect_heading_level(
        cls, style_name: str, docx: DocxDocument, pPr=None,
    ) -> int | None:
        """从样式名或段落属性中检测标题层级。

        采用多层检测策略，按优先级依次尝试：
        1. 样式 ID 直接匹配多语言标题关键词并提取数字
        2. 样式 ID 含关键词但无数字时，查样式表取层级
        3. 样式 ID 不含关键词时，查样式表按样式名匹配关键词（修复核心：style_id 可能
           是 "1"/"a1" 等不含关键词的值，需通过样式表获取真实名称如 "Heading 1"）
        4. w:outlineLvl 大纲级别兜底（OOXML 标准机制）

        Args:
            style_name: w:pStyle / w:val 属性值（样式 ID）
            docx: python-docx Document 对象，用于样式表二次确认
            pPr: w:pPr XML 元素，用于检测 w:outlineLvl

        Returns:
            标题层级（1-9），非标题样式返回 None。
        """
        if not style_name:
            return cls._detect_by_outline_level(pPr)

        name_lower = style_name.lower()

        # ── 策略 1：样式 ID 直接匹配关键词 ──
        matched_keyword = None
        for kw in cls.HEADING_KEYWORDS:
            if kw in name_lower:
                matched_keyword = kw
                break

        if matched_keyword is not None:
            # 去掉关键词后提取尾部连续数字
            remaining = name_lower.replace(matched_keyword, "").strip()
            digits = ""
            for ch in reversed(remaining):
                if ch.isdigit():
                    digits = ch + digits
                else:
                    break
            if digits:
                return int(digits)

            # 有关键词但无数字 → 查样式表取层级
            level = cls._lookup_style_heading_level(style_name, docx)
            if level is not None:
                return level
            return 1  # 默认 level 1

        # ── 策略 2：样式 ID 不含关键词 → 查样式表按样式名匹配 ──
        level = cls._lookup_style_heading_level(style_name, docx)
        if level is not None:
            return level

        # ── 策略 3：w:outlineLvl 兜底 ──
        return cls._detect_by_outline_level(pPr)

    @staticmethod
    def _detect_by_outline_level(pPr) -> int | None:
        """从 w:outlineLvl 检测大纲级别（OOXML 标准标题机制）。

        outlineLvl 值从 0 开始（0 = Level 1 标题），返回时 +1 对齐 heading_level。
        部分文档生成器使用 w:outlineLvl 而非 Heading 样式来标记标题。
        """
        if pPr is None:
            return None
        outline_lvl = pPr.find(qn("w:outlineLvl"))
        if outline_lvl is not None:
            try:
                lvl = int(outline_lvl.get(qn("w:val"), "0"))
                return lvl + 1  # outlineLvl 0-based → heading_level 1-based
            except (ValueError, TypeError):
                pass
        return None

    @classmethod
    def _lookup_style_heading_level(
        cls, style_id: str, docx: DocxDocument,
    ) -> int | None:
        """在 docx.styles 中按 style_id 查找样式，检查其名称是否包含标题关键词。

        与样式 ID 匹配逻辑解耦：即使样式 ID 是 "1" / "a1" 等不含关键词的值，
        也能通过样式表获取真实名称（如 "Heading 1" / "标题 1"）来判定。

        Args:
            style_id: 样式 ID（w:pStyle / w:val 的值）
            docx: python-docx Document 对象

        Returns:
            标题层级（1-9），非标题样式返回 None。
        """
        try:
            for style in docx.styles:
                if style.style_id == style_id and style.type == WD_STYLE_TYPE.PARAGRAPH:
                    if style.name:
                        s_lower = style.name.lower()
                        for kw in cls.HEADING_KEYWORDS:
                            if kw in s_lower:
                                # 从样式名称中提取层级数字
                                remain = s_lower.replace(kw, "").strip()
                                d = ""
                                for ch in reversed(remain):
                                    if ch.isdigit():
                                        d = ch + d
                                    else:
                                        break
                                if d:
                                    return int(d)
                                return 1  # 有关键词但无数字，默认 level 1
                    break
        except Exception:
            pass
        return None

    # ── 段落处理 ────────────────────────────────────────────────

    def _process_paragraph(
        self, p_el, docx: DocxDocument, state: "_DocxParseState"
    ) -> None:
        """处理一w:p 元素（段落）
        w:p 的直接子元素（w:r w:hyperlink）顺序遍历，
        提取内联图片、超链接和文本，保持原文顺序        """
        # 获取样式名称和列表标记。
        style_name = ""
        is_list = False
        pPr = p_el.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                style_name = pStyle.get(qn("w:val"), "")
                is_list = style_name.lower().startswith("list")
            is_list = is_list or pPr.find(qn("w:numPr")) is not None

        # 按直接子元素顺序遍历，提取文本和资源
        text_parts: list[str] = []
        has_content = False  # 是否有实际内容（文本、图片或链接）
        in_field = False      # 是否在 w:fldChar begin..end 范围内
        field_has_asset = False     # 字段指令已创建 Asset
        pending_placeholder = ""    # 待追加到显示文字后的占位符
        for child in p_el:
            child_tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if child_tag == "r":
                # ── 文本运行（w:r）──
                # 跟踪 w:fldChar 状态
                fld_char_type = self._get_fld_char_type(child)
                if fld_char_type == "begin":
                    in_field = True
                    continue
                if fld_char_type == "separate":
                    continue
                if fld_char_type == "end":
                    in_field = False
                    field_has_asset = False
                    pending_placeholder = ""
                    continue

                # 先检查内联图片/视频节点（w:drawing），仅追加占位符
                drawing_rIds = self._extract_drawing_rIds(child)
                for rId in drawing_rIds:
                    asset = self._resolve_image_asset(rId, docx, state)
                    if asset is not None:
                        placeholder = state.next_placeholder(asset.asset_type.value)
                        state.track_asset(AssetData(
                            placeholder=placeholder, asset_id=asset.asset_id,
                        ))
                        text_parts.append(placeholder)
                        has_content = True

                # 检查 w:instrText 字段指令（企业微信微盘等嵌入文件），文字被占位符替换
                instr_url, instr_filename = self._parse_field_instruction(child)
                if instr_url and instr_filename:
                    asset_type = classify_link_text(instr_filename)
                    placeholder = state.next_placeholder(asset_type.value)
                    asset = Asset(
                        doc_id=state.doc_id,
                        asset_type=asset_type,
                        original_uri=instr_url,
                        display_text=instr_filename,
                        status=AssetStatus.ready,
                        metadata={
                            "source": "field_instruction",
                        },
                    )
                    state.assets.append(asset)
                    state.track_asset(AssetData(
                        placeholder=placeholder, asset_id=asset.asset_id,
                    ))
                    has_content = True
                    field_has_asset = True
                    pending_placeholder = placeholder
                    continue

                # 字段显示文字：被占位符替换，不保留原文
                if in_field and field_has_asset:
                    text_parts.append(pending_placeholder)
                    field_has_asset = False
                    has_content = True
                    continue

                # 提取文本节点。
                for t_node in child.iter(qn("w:t")):
                    if t_node.text:
                        text_parts.append(t_node.text)
                        has_content = True

            elif child_tag == "hyperlink":
                # ── 超链接（w:hyperlink）──
                link_text, url = self._extract_hyperlink(child, docx)
                if url:
                    # 按链接文字后缀分类（非 URL），兜底为 web_link
                    asset_type = classify_link_text(link_text)
                    asset = Asset(
                        doc_id=state.doc_id,
                        asset_type=asset_type,
                        original_uri=url,
                        display_text=link_text,
                        status=AssetStatus.ready,
                        metadata={},
                    )
                    state.assets.append(asset)
                    placeholder = state.next_placeholder(asset_type.value)
                    # 链接文字被占位符替换
                    text_parts.append(placeholder)
                    has_content = True
                    state.track_asset(AssetData(
                        placeholder=placeholder, asset_id=asset.asset_id,
                    ))

        text = "".join(text_parts).strip()

        # 检查不受支持的嵌入对象
        has_unknown_object = self._has_unsupported_object(p_el)

        # 空段落没有文本、图片或链接时跳过。
        if not text and not has_content:
            if has_unknown_object:
                state.add_unknown("[不支持的嵌入 DOCX 对象]")
            return

        # 根据样式判断元素类型
        # 标题检测优先：即使段落带有 w:numPr（自动编号），只要样式为
        # Heading 也应识别为标题，避免被误判为列表项。
        heading_match = self._detect_heading_level(style_name, docx, pPr)

        if heading_match is not None:
            state.add_title(text, heading_match)
        elif is_list:
            state.add_to_list_buffer(text)
        elif state._list_buffer:
            # 列表缓冲区非空时，即使当前段落不带 w:numPr，
            # 也视为同一列表的延续，避免被无编号的中间段落打断。
            state.add_to_list_buffer(text)
        else:
            state.add_paragraph(text)

        if has_unknown_object:
            state.add_unknown("[不支持的嵌入 DOCX 对象]")

    # ── 超链接提──────────────────────────────────────────────

    def _extract_hyperlink(self, hyperlink_el, docx: DocxDocument) -> tuple[str, str]:
        """w:hyperlink 元素中提取显示文字和目标 URL
        Args:
            hyperlink_el: w:hyperlink XML 元素            docx: python-docx Document 对象
        Returns:
            (显示文字, 目标URL) 元组，不存在时为空字符串        """
        # 提取显示文字：遍历内部所w:t 节点
        texts = []
        for t_node in hyperlink_el.iter(qn("w:t")):
            if t_node.text:
                texts.append(t_node.text)
        display_text = "".join(texts)

        # 提取目标 URL：通过 r:id docx.part.rels rel.target_ref
        rId = hyperlink_el.get(qn("r:id"))
        url = ""
        if rId:
            try:
                rel = docx.part.rels[rId]
                url = rel.target_ref
            except (KeyError, AttributeError):
                pass

        return display_text, url

    # ── 链接分类 ────────────────────────────────────────────────

    # ── 字段字符检测 ──────────────────────────────────────────

    @staticmethod
    def _get_fld_char_type(r_el) -> str | None:
        """从 w:r 中提取 w:fldChar 的 fldCharType 属性值。

        Returns:
            "begin" / "separate" / "end"，不存在时返回 None。
        """
        fld = r_el.find(qn("w:fldChar"))
        if fld is not None:
            return fld.get(qn("w:fldCharType"))
        return None

    # ── 字段指令解析（企业微信微盘等嵌入文件）─────────────────

    @staticmethod
    def _parse_field_instruction(r_el) -> tuple[str, str]:
        """从 w:r 的 w:instrText 中提取 URL 和文件名。

        支持多种字段格式：
        - WeDrive: \\tdfu https://... \\tdfn file.mov
        - HYPERLINK: HYPERLINK \"https://...\"
        - INCLUDEPICTURE: INCLUDEPICTURE \"https://...\"
        - 通用：提取第一个 https?:// URL 作为链接

        Returns:
            (url, filename) 元组，未找到时返回 ("", "")。
        """
        for instr_node in r_el.iter(qn("w:instrText")):
            text = instr_node.text or ""

            # 策略1：WeDrive 嵌入文件（\\tdfu URL + \\tdfn 文件名）
            url_match = re.search(r"\\tdfu\s+(\S+)", text)
            fn_match = re.search(r"\\tdfn\s+(\S+)", text)
            if url_match and fn_match:
                return url_match.group(1), fn_match.group(1)

            # 策略2：HYPERLINK / INCLUDEPICTURE 等标准字段
            m = re.search(r'''(?:HYPERLINK|INCLUDEPICTURE)\s+"(https?://[^"]+)"''', text, re.IGNORECASE)
            if m:
                url = m.group(1)
                filename = url.rsplit("/", 1)[-1].split("?", 1)[0] or "file"
                return url, filename

            # 策略3：通用兜底 — 提取任意 https?:// URL
            m = re.search(r"(https?://\S+)", text)
            if m:
                url = m.group(0)
                filename = url.rsplit("/", 1)[-1].split("?", 1)[0] or "file"
                return url, filename
        return "", ""

    # ── 内联图片处理 ────────────────────────────────────────────

    @staticmethod
    def _extract_drawing_rIds(r_el) -> list[str]:
        """w:r 元素中提取所有内联图片的 r:embed rId
        遍历 w:drawing a:blip 提取 r:embed 属性
        Args:
            r_el: w:r XML 元素
        Returns:
            r:embed 属性值列表        """
        rIds: list[str] = []
        for drawing in r_el.iter(qn("w:drawing")):
            for blip in drawing.iter():
                blip_tag = blip.tag.split("}")[-1] if "}" in blip.tag else blip.tag
                if blip_tag == "blip":
                    embed = blip.get(qn("r:embed"))
                    if embed:
                        rIds.append(embed)
        return rIds

    def _resolve_image_asset(
        self, rId: str, docx: DocxDocument, state: "_DocxParseState"
    ) -> Asset | None:
        """通过 rId asset map 中查找对应的图片 Asset
        Args:
            rId: w:drawing a:blip r:embed 的值            docx: python-docx Document 对象            state: 当前解析状态（_image_asset_map）
        Returns:
            匹配Asset，未找到时返None        """
        try:
            rel = docx.part.rels[rId]
            target = rel.target_ref  # 格式为 media/image1.png，不含 word/ 前缀。
            return state._image_asset_map.get(target)
        except (KeyError, AttributeError):
            return None

    # ── 表格处理 ────────────────────────────────────────────────

    def _process_table(
        self, tbl_el, docx: DocxDocument, state: "_DocxParseState"
    ) -> None:
        """处理一w:tbl 元素（表格）
        表格的出现意味着列表上下文结束，先刷新列表缓冲区，
        再提取行列结构，处理合并单元格（含垂直合并），
        同时提取单元格内的图片和超链接资源        """
        state._flush_list_buffer()
        # 每格(text, asset_data) 元组
        rows_data: list[list[tuple[str, list[AssetData]]]] = []
        headers: list[tuple[str, list[AssetData]]] = []
        vertical_merges: dict[int, tuple[str, list[AssetData]]] = {}

        for i, tr in enumerate(tbl_el.findall(qn("w:tr"))):
            cells: list[tuple[str, list[AssetData]]] = []
            col_idx = 0
            for tc in tr.findall(qn("w:tc")):
                # 提取单元格的文本和资源信息
                cell_text_parts: list[str] = []
                cell_asset_data: list[AssetData] = []

                # w:tc 的直接子元素w:p（段落），遍历每w:p
                for tc_child in tc:
                    tc_tag = tc_child.tag.split("}")[-1] if "}" in tc_child.tag else tc_child.tag
                    if tc_tag != "p":
                        continue

                    # 复用段落子元素遍历模板
                    para_text_parts: list[str] = []
                    for p_child in tc_child:
                        p_tag = p_child.tag.split("}")[-1] if "}" in p_child.tag else p_child.tag

                        if p_tag == "r":
                            # 内联图片/视频，仅追加占位符
                            drawing_rIds = self._extract_drawing_rIds(p_child)
                            for rId in drawing_rIds:
                                asset = self._resolve_image_asset(rId, docx, state)
                                if asset is not None:
                                    placeholder = state.next_placeholder(asset.asset_type.value)
                                    cell_asset_data.append(AssetData(
                                        placeholder=placeholder, asset_id=asset.asset_id,
                                    ))
                                    para_text_parts.append(placeholder)
                            # 文本
                            for t_node in p_child.iter(qn("w:t")):
                                if t_node.text:
                                    para_text_parts.append(t_node.text)

                        elif p_tag == "hyperlink":
                            link_text, url = self._extract_hyperlink(p_child, docx)
                            if url:
                                asset_type = classify_link_text(link_text)
                                asset = Asset(
                                    doc_id=state.doc_id,
                                    asset_type=asset_type,
                                    original_uri=url,
                                    display_text=link_text,
                                    status=AssetStatus.ready,
                                    metadata={},
                                )
                                state.assets.append(asset)
                                placeholder = state.next_placeholder(asset_type.value)
                                para_text_parts.append(placeholder)
                                cell_asset_data.append(AssetData(
                                    placeholder=placeholder, asset_id=asset.asset_id,
                                ))
                    if para_text_parts:
                        cell_text_parts.append("".join(para_text_parts))

                # 多个 w:p 之间用换行符分隔
                cell_text = "\n".join(cell_text_parts) if cell_text_parts else ""

                # 获取合并单元格属性。
                span = 1
                vmerge_val = None
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is not None:
                    grid_span = tcPr.find(qn("w:gridSpan"))
                    if grid_span is not None:
                        try:
                            span = max(1, int(grid_span.get(qn("w:val"), "1")))
                        except ValueError:
                            span = 1
                    vmerge = tcPr.find(qn("w:vMerge"))
                    if vmerge is not None:
                        vmerge_val = vmerge.get(qn("w:val")) or "continue"

                # 处理垂直合并单元格（asset_data 传递）
                if vmerge_val == "continue":
                    merged = vertical_merges.get(col_idx, ("", []))
                    cell_text = merged[0] or cell_text
                    cell_asset_data = merged[1] or cell_asset_data
                elif vmerge_val == "restart":
                    for offset in range(span):
                        vertical_merges[col_idx + offset] = (cell_text, list(cell_asset_data))

                for offset in range(span):
                    if vmerge_val == "continue":
                        merged = vertical_merges.get(col_idx + offset, ("", []))
                        cells.append((merged[0] or cell_text, list(merged[1] or cell_asset_data)))
                    else:
                        cells.append((cell_text, list(cell_asset_data)))
                col_idx += span

            if i == 0:
                headers = cells
            else:
                rows_data.append(cells)

        if not rows_data:
            return

        # 构建 structured_data
        all_asset_data: list[AssetData] = []
        structured: dict[str, Any] = {
            "table": {
                "caption": "",
                "headers": [{"text": h[0]} for h in headers],
                "rows": [
                    {
                        "cells": [{"text": cell[0]} for cell in row]
                    }
                    for row in rows_data
                ],
            }
        }
        # 表格级资源引用按 asset_id 去重，避免合并单元格造成重复引用。
        seen_asset_ids: set[str] = set()
        for h in headers:
            for ad in h[1]:
                if ad.asset_id not in seen_asset_ids:
                    seen_asset_ids.add(ad.asset_id)
                    all_asset_data.append(ad)
        for row in rows_data:
            for cell in row:
                for ad in cell[1]:
                    if ad.asset_id not in seen_asset_ids:
                        seen_asset_ids.add(ad.asset_id)
                        all_asset_data.append(ad)

        # 构建纯文本表示（仅使用文本部分）
        flat_parts = []
        if headers:
            flat_parts.append(" | ".join(h[0] for h in headers))
        for row in rows_data:
            flat_parts.append(" | ".join(cell[0] for cell in row))

        element_kwargs: dict[str, Any] = {
            "doc_id": state.doc_id,
            "doc_version": state.doc_version,
            "parent_element_id": None,
            "sequence_order": state._next_seq(),
            "element_type": ElementType.table,
            "text": "\n".join(flat_parts),
            "structured_data": structured,
            "asset_data": all_asset_data,
            "source_location": SourceLocation(section_path=list(state._section_path)),
        }
        el = ParsedElement(**element_kwargs)
        state.elements.append(el)

    # ── 图片预提──────────────────────────────────────────────

    def _build_image_asset_map(
        self, content: bytes, doc: Document, state: "_DocxParseState"
    ) -> None:
        """docx 归档 word/media/ 目录预提取所有嵌入媒体文件。

        图片和视频统一按扩展名分类存入 state._image_asset_map 和 state.assets。
        通过 content 参数直接接收原始字节，避免依赖 doc.metadata。

        Args:
            content: 文档原始字节（zip 格式）
            doc: 文档对象
            state: 当前解析状态
        """
        if isinstance(content, str):
            content = content.encode("utf-8")
        zip_source = content

        try:
            with zipfile.ZipFile(
                io.BytesIO(zip_source) if isinstance(zip_source, bytes) else zip_source
            ) as zf:
                media_files = [
                    name for name in zf.namelist()
                    if name.startswith("word/media/") and not name.endswith("/")
                ]
                for idx, name in enumerate(media_files):
                    data = zf.read(name)
                    content_hash = hashlib.sha256(data).hexdigest()
                    ext = name.rsplit(".", 1)[-1].lower() if "." in name else "bin"
                    filename = name.split("/")[-1]

                    # 根据扩展名区分嵌入图片和视频
                    is_video = ext in {"mov", "mp4", "webm", "m4v"}
                    asset_type = AssetType.video if is_video else AssetType.image

                    # 同时保存 ZIP 完整路径与关系文件使用的短路径。
                    short_path = name.replace("word/", "", 1)  # media/image1.png

                    asset = Asset(
                        doc_id=state.doc_id,
                        element_id="",
                        asset_type=asset_type,
                        original_uri="",                # 嵌入类型无外部来源
                        content_hash=f"sha256:{content_hash}",
                        status=AssetStatus.ready,
                        storage_uri=None,
                        extracted_text=None,
                        metadata={},
                    )
                    object.__setattr__(asset, "_data", data)

                    # key 存储
                    state._image_asset_map[name] = asset       # word/media/image1.png
                    state._image_asset_map[short_path] = asset  # media/image1.png
                    state.assets.append(asset)

        except (zipfile.BadZipFile, KeyError) as exc:
            logger.warning("docx 提取图片失败: %s", exc)

    # ── 视频提取 ────────────────────────────────────────────────

    # ── 资源关联回填 ────────────────────────────────────────────

    @staticmethod
    def _link_assets_to_elements(
        elements: list[ParsedElement],
        assets: list[Asset],
    ) -> None:
        """在元素获得 ID 后，通过 asset_id 回填 Asset.element_id。"""
        elements_by_asset_id = {
            ad.asset_id: el.element_id
            for el in elements
            for ad in el.asset_data
        }
        for asset in assets:
            if not asset.element_id:
                asset.element_id = elements_by_asset_id.get(asset.asset_id, "")

    # ── 不支持对象检──────────────────────────────────────────

    def _has_unsupported_object(self, p_el) -> bool:
        """检查段落中是否包含不受支持的嵌入对象（OLE 等）。"""
        unsupported_tags = {"object", "oleobject", "control"}
        for node in p_el.iter():
            local = node.tag.split("}")[-1].lower() if "}" in node.tag else node.tag.lower()
            if local in unsupported_tags:
                return True
        return False


# ── 内部解析状────────────────────────────────────────────────


@dataclass
class _DocxParseState(_BaseParseState):
    """DOCX 解析过程中的可变状态，在遍body 元素时逐步构建元素
    继承 _BaseParseState doc_id、doc_version、elements、_seq、_section_path
    _next_seq() 方法    扩展资源跟踪字段，支持内联图片和超链接的关联    """

    _current_list_id: str | None = None
    _list_buffer: list[str] = field(default_factory=list)  # 连续列表项文本缓冲区
    _tracked_assets: list[AssetData] = field(default_factory=list)
    _link_urls: list[str] = field(default_factory=list)
    _image_asset_map: dict[str, Asset] = field(default_factory=dict)
    _counters: dict[str, int] = field(default_factory=dict)
    assets: list[Asset] = field(default_factory=list)

    # ── 资源跟踪 ────────────────────────────────────────────────

    def track_asset(self, asset_data: AssetData) -> None:
        """记录当前上下文关联的资源信息，在添加段落时消费。"""
        self._tracked_assets.append(asset_data)

    def track_link_url(self, url: str) -> None:
        """记录当前段落的链接 URL，在添加段落时消费。"""
        self._link_urls.append(url)

    def consume_tracked_assets(self) -> list[AssetData]:
        """消费并清空资源跟踪列表。"""
        result = list(self._tracked_assets)
        self._tracked_assets = []
        return result

    # 资源类型 → 占位符前缀映射
    _PLACEHOLDER_PREFIX = {
        "video_link": "video",
        "video": "video",
        "image_link": "image",
        "image": "image",
        "document_link": "doc",
        "web_link": "web",
    }

    def next_placeholder(self, asset_type: str) -> str:
        """生成递增占位符，如 {{video:1}}, {{image:2}}, {{doc:3}}, {{web:4}}。

        与 md 解析器行为一致，每种类型独立计数。
        """
        prefix = self._PLACEHOLDER_PREFIX.get(asset_type, "res")
        self._counters[prefix] = self._counters.get(prefix, 0) + 1
        return f"{{{{{prefix}:{self._counters[prefix]}}}}}"

    def consume_link_urls(self) -> list[str]:
        """消费并清空链接 URL 列表。"""
        result = list(self._link_urls)
        self._link_urls = []
        return result

    # ── 元素累积 ────────────────────────────────────────────────

    def flush_elements(self) -> list[ParsedElement]:
        """完成解析：先刷新列表缓冲区，再返回全部累积元素。"""
        self._flush_list_buffer()
        return self.elements

    def add_title(self, text: str, level: int) -> None:
        """添加标题元素并按层级更新 section_path。

        标题出现时先刷新列表缓冲区，再清空资源跟踪。
        """
        if not text:
            return
        self._flush_list_buffer()
        self._current_list_id = None
        self.consume_tracked_assets()  # 丢弃
        self.consume_link_urls()       # 丢弃
        while len(self._section_path) >= level:
            self._section_path.pop()
        self._section_path.append(text)
        self.elements.append(
            ParsedElement(
                doc_id=self.doc_id,
                doc_version=self.doc_version,
                sequence_order=self._next_seq(),
                element_type=ElementType.title,
                text=text,
                source_location=SourceLocation(section_path=list(self._section_path)),
                metadata={"heading_level": level},
            )
        )

    def add_paragraph(self, text: str) -> None:
        """添加普通段落元素，并消费跟踪的资源。

        段落的到来意味着列表上下文结束，先刷新列表缓冲区。
        """
        self._flush_list_buffer()
        self._current_list_id = None
        asset_data = self.consume_tracked_assets()
        self.consume_link_urls()  # 已废弃，link_urls 不再写入 metadata
        element = ParsedElement(
            doc_id=self.doc_id,
            doc_version=self.doc_version,
            sequence_order=self._next_seq(),
            element_type=ElementType.paragraph,
            text=text,
            asset_data=asset_data,
            source_location=SourceLocation(section_path=list(self._section_path)),
        )
        self.elements.append(element)

    def add_to_list_buffer(self, text: str) -> None:
        """将列表项文本追加到缓冲区，不清空资源跟踪。

        多个连续列表项缓冲后，在遇到下一个非列表元素或解析结束时
        由 _flush_list_buffer() 合并为一个 paragraph 整体输出。
        """
        self.consume_tracked_assets()  # 列表项不关联资源
        self.consume_link_urls()
        self._list_buffer.append(text)

    def _flush_list_buffer(self) -> None:
        """将缓冲区中的连续列表项合并为一个 paragraph 元素并清空缓冲区。"""
        if not self._list_buffer:
            return
        merged_text = "\n".join(self._list_buffer)
        self._list_buffer.clear()
        self._current_list_id = None
        self.elements.append(
            ParsedElement(
                doc_id=self.doc_id,
                doc_version=self.doc_version,
                sequence_order=self._next_seq(),
                element_type=ElementType.paragraph,
                text=merged_text,
                source_location=SourceLocation(section_path=list(self._section_path)),
            )
        )

    def add_list_item(self, text: str) -> None:
        """向后兼容的列表项添加方法，内部委托给 add_to_list_buffer。"""
        self.add_to_list_buffer(text)

    def add_unknown(self, text: str) -> None:
        """添加未知类型元素，例如不支持的内嵌对象。"""
        self.elements.append(
            ParsedElement(
                doc_id=self.doc_id,
                doc_version=self.doc_version,
                sequence_order=self._next_seq(),
                element_type=ElementType.unknown,
                text=text,
                source_location=SourceLocation(section_path=list(self._section_path)),
            )
        )
