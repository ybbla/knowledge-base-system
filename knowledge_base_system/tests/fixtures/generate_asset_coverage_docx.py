"""生成覆盖所有 ElementType + AssetType 的测试 DOCX，并输出解析结果概览。

运行方式：
    cd knowledge_base_system
    python tests/fixtures/generate_asset_coverage_docx.py

用途：验证 docx_parser 对各种元素和资源类型的处理能力。
"""

import base64
import io
import os
import sys
import tempfile

# 修复 Windows GBK 终端 emoji 编码问题
import codecs

sys.stdout = codecs.getwriter("utf-8")(sys.stdout.buffer, "strict")

from docx import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# 确保能导入项目模块
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", ".."))

from app.core.models import AssetType, Document, ElementType
from parsers.docx_parser import DocxParser


# ── 工具函数 ──────────────────────────────────────────────────────────


def _add_hyperlink(paragraph, url: str, text: str, docx_part) -> None:
    """在段落中追加一个超链接（不改变已有内容）。"""
    rel = docx_part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def _add_unsupported_object(paragraph) -> None:
    """追加一个不支持的 OLE 嵌入对象。"""
    obj = OxmlElement("w:object")
    paragraph._p.append(obj)


def _make_1x1_png() -> bytes:
    """最小 1x1 PNG（base64 解码）。"""
    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
    )


# ── 构建测试文档 ──────────────────────────────────────────────────────


def build_test_docx() -> bytes:
    """创建覆盖全部 ElementType + AssetType 的测试 DOCX。

    内容结构（自上而下）：
    1.  Heading 1              → ElementType.title
    2.  Paragraph              → ElementType.paragraph
    3.  List Bullet × 2        → ElementType.list + paragraph(items)
    4.  Table (含合并单元格)    → ElementType.table
    5.  Inline Image           → AssetType.image (内嵌图片)
    6.  Image Link             → AssetType.image_link (超链接指向图片)
    7.  Video Link             → AssetType.video_link (超链接指向视频)
    8.  Document Link          → AssetType.document_link (超链接指向 PDF)
    9.  Video in Text          → AssetType.video_link (正文中含视频 URL)
    10. Unsupported Object     → ElementType.unknown
    """
    docx = DocxDocument()
    png_bytes = _make_1x1_png()
    image_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    image_file.write(png_bytes)
    image_file.close()

    try:
        # ── 1. 标题 ─────────────────────────────────────────
        docx.add_heading("知识库系统设计文档", level=1)
        docx.add_heading("架构概述", level=2)

        # ── 2. 段落 ─────────────────────────────────────────
        docx.add_paragraph(
            "本文档描述了知识库系统的整体架构设计，包括解析层、索引层和检索层。"
        )

        # ── 3. 列表 ─────────────────────────────────────────
        docx.add_paragraph("解析器支持多种格式", style="List Bullet")
        docx.add_paragraph("索引使用 Milvus + BM25 双路", style="List Bullet")
        docx.add_paragraph("检索采用 RRF 融合 + LLM Rerank", style="List Bullet")

        # ── 4. 表格（含合并单元格）─────────────────────────
        table = docx.add_table(rows=4, cols=3)
        table.style = "Table Grid"
        table.cell(0, 0).text = "模块"
        table.cell(0, 1).text = "技术栈"
        table.cell(0, 2).text = "说明"
        table.cell(1, 0).text = "解析层"
        table.cell(1, 1).text = "python-docx/PyMuPDF"
        table.cell(1, 2).text = "多格式文档解析"
        table.cell(2, 0).text = "索引层"
        table.cell(2, 1).text = "Milvus 2.5"
        table.cell(2, 2).text = "向量 + BM25 双路"
        table.cell(3, 0).text = "合并示例"
        table.cell(3, 0).merge(table.cell(3, 1))
        table.cell(3, 2).text = "跨两列合并"

        # ── 5. 内嵌图片 ─────────────────────────────────────
        p_img = docx.add_paragraph()
        run_img = p_img.add_run()
        run_img.add_picture(image_file.name)
        p_img.add_run(" 上图：系统架构图")

        # ── 6. 图片链接 ─────────────────────────────────────
        p_img_link = docx.add_paragraph()
        p_img_link.add_run("系统架构图（外链）：")
        _add_hyperlink(
            p_img_link,
            "https://example.com/architecture.png",
            "architecture.png",
            docx.part,
        )

        # ── 7. 视频链接 ─────────────────────────────────────
        p_video = docx.add_paragraph()
        p_video.add_run("演示视频：")
        _add_hyperlink(
            p_video,
            "https://www.youtube.com/watch?v=dQw4w9WgXcQ",
            "系统演示",
            docx.part,
        )

        # ── 8. 文档链接 ─────────────────────────────────────
        p_doc = docx.add_paragraph()
        p_doc.add_run("详细设计文档：")
        _add_hyperlink(
            p_doc,
            "https://example.com/design-spec.pdf",
            "design-spec.pdf",
            docx.part,
        )

        # ── 9. 正文中的视频 URL ─────────────────────────────
        docx.add_paragraph(
            "也可通过以下链接观看：https://www.bilibili.com/video/BV1xx411c7mD/ 回放录像。"
        )

        # ── 10. 不支持对象 ──────────────────────────────────
        p_obj = docx.add_paragraph("以下为嵌入对象：")
        _add_unsupported_object(p_obj)

        buf = io.BytesIO()
        docx.save(buf)
        return buf.getvalue()
    finally:
        os.unlink(image_file.name)


# ── 解析与概览输出 ────────────────────────────────────────────────────


def print_separator(title: str) -> None:
    print()
    print("=" * 72)
    print(f"  {title}")
    print("=" * 72)


def print_parse_summary(content: bytes) -> None:
    """解析 DOCX 并打印结构化的结果概览。"""
    parser = DocxParser()
    doc = Document(
        title="Asset Coverage Test",
        source_type="docx",
        source_uri="memory://asset-coverage-test",
        metadata={"raw_content": content},
    )
    result = parser.parse(doc, doc.metadata["raw_content"])

    # ── 按类型统计元素 ──
    print_separator("ElementType 覆盖统计")
    type_counts: dict[ElementType, int] = {}
    for el in result.elements:
        type_counts[el.element_type] = type_counts.get(el.element_type, 0) + 1

    all_types = list(ElementType)
    for et in all_types:
        count = type_counts.get(et, 0)
        status = "✅" if count > 0 else "❌ 未覆盖"
        print(f"  {et.value:15s} : {count:3d} 个  {status}")

    # ── 按类型统计 Asset ──
    print_separator("AssetType 覆盖统计")
    asset_type_counts: dict[AssetType, int] = {}
    for a in result.assets:
        asset_type_counts[a.asset_type] = asset_type_counts.get(a.asset_type, 0) + 1

    for at in list(AssetType):
        count = asset_type_counts.get(at, 0)
        status = "✅" if count > 0 else "❌ 未覆盖"
        print(f"  {at.value:16s} : {count:3d} 个  {status}")

    # ── 逐元素详情 ──
    print_separator("逐元素详情")
    for i, el in enumerate(result.elements):
        text_preview = (el.text or "")[:80].replace("\n", "\\n")
        type_tag = el.element_type.value

        # asset_data 摘要
        ad_summary = ""
        if el.asset_data:
            types = [ad.type for ad in el.asset_data]
            ad_summary = f" | assets: {types}"

        # metadata 摘要
        meta_summary = ""
        if el.metadata:
            keys = list(el.metadata.keys())
            meta_summary = f" | meta: {keys}"

        # structured_data 摘要
        sd_summary = ""
        if el.structured_data:
            sd_keys = list(el.structured_data.keys())
            sd_summary = f" | structured: {sd_keys}"

        heading_info = ""
        if el.element_type == ElementType.title:
            heading_info = f" L{el.metadata.get('heading_level', '?')}"

        parent_info = ""
        if el.parent_element_id:
            parent_info = f" parent={el.parent_element_id}"

        print(
            f"  [{i:02d}] {type_tag:12s}{heading_info} "
            f"\"{text_preview}\"{ad_summary}{meta_summary}{sd_summary}{parent_info}"
        )

    # ── Asset 详情 ──
    print_separator("Asset 详情")
    for i, a in enumerate(result.assets):
        source_info = ""
        if a.element_id:
            source_info = f" -> element: {a.element_id}"
        mime_info = a.mime_type or "?"
        print(
            f"  [{i:02d}] {a.asset_type.value:16s} "
            f"mime={mime_info:24s} "
            f"uri={a.original_uri}{source_info}"
        )

    # ── 覆盖总结 ──
    print_separator("覆盖总结")
    covered_et = {et for et in all_types if type_counts.get(et, 0) > 0}
    missing_et = set(all_types) - covered_et
    covered_at = {at for at in AssetType if asset_type_counts.get(at, 0) > 0}
    missing_at = set(AssetType) - covered_at

    print(f"  ElementType: {len(covered_et)}/{len(all_types)} 已覆盖")
    if missing_et:
        print(f"    缺失: {[e.value for e in missing_et]}")
    print(f"  AssetType:   {len(covered_at)}/{len(AssetType)} 已覆盖")
    if missing_at:
        print(f"    缺失: {[a.value for a in missing_at]}")

    print()


# ── 入口 ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    content = build_test_docx()
    print(f"测试 DOCX 已生成，大小: {len(content):,} bytes")
    print_parse_summary(content)
