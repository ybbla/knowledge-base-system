"""Phase 2 端到端 API 测试脚本，并输出 Markdown/JSON 报告。

启动方式：
    1. 安装依赖：pip install -r requirements.txt
    2. 如需真实语义抽取和向量检索结果，配置 VOLCENGINE_API_KEY。
    3. 在项目根目录运行：
       python knowledge_base_system/tests/e2e/e2e_phase2_markdown_docx.py
    4. 未配置 VOLCENGINE_API_KEY 时可加 --allow-zero-chunks，仅验证 API 链路：
       python knowledge_base_system/tests/e2e/e2e_phase2_markdown_docx.py --allow-zero-chunks

执行流程：
    1. 使用 TestClient 加载真实 FastAPI app，不需要单独启动 uvicorn。
    2. 对 Markdown 样例和自动生成的 DOCX 样例分别执行：
       GET /health -> POST /upload -> POST /ingest -> GET /ingest/{job_id} -> POST /search。
    3. 每一步记录请求摘要、响应内容、校验项、耗时和异常信息。

结果保存：
    - 测试输入样例：knowledge_base_system/tests/fixtures/
    - 生成报告目录：knowledge_base_system/tests/results/e2e/
    - 每次运行生成 phase2_e2e_YYYYMMDD_HHMMSS.md 和 .json。
"""

from __future__ import annotations

import argparse
import base64
import json
import mimetypes
import os
import sys
import tempfile
import time
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from fastapi.testclient import TestClient


SCRIPT_DIR = Path(__file__).resolve().parent
TESTS_DIR = SCRIPT_DIR.parent
PACKAGE_ROOT = TESTS_DIR.parent
REPO_ROOT = PACKAGE_ROOT.parent
INPUT_DIR = TESTS_DIR / "fixtures" / "simulated_inputs"
GENERATED_INPUT_DIR = TESTS_DIR / "fixtures" / "generated_inputs"
RESULTS_DIR = TESTS_DIR / "results" / "e2e"
DEFAULT_MARKDOWN_INPUT = INPUT_DIR / "product_manual_source.md"
DEFAULT_CATEGORY = "phase2-e2e"

if str(PACKAGE_ROOT) not in sys.path:
    sys.path.insert(0, str(PACKAGE_ROOT))

mimetypes.add_type("text/markdown", ".md")
mimetypes.add_type("text/markdown", ".markdown")
mimetypes.add_type(
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".docx",
)


@dataclass
class Check:
    name: str
    passed: bool
    detail: str = ""


@dataclass
class Step:
    name: str
    ok: bool
    elapsed_seconds: float = 0.0
    request: dict[str, Any] = field(default_factory=dict)
    response: Any = None
    checks: list[Check] = field(default_factory=list)
    error: str | None = None


@dataclass
class CaseResult:
    name: str
    input_path: str
    source_type: str
    steps: list[Step]

    @property
    def ok(self) -> bool:
        return all(step.ok for step in self.steps)


def _jsonable(value: Any) -> Any:
    try:
        json.dumps(value, ensure_ascii=False)
        return value
    except TypeError:
        return str(value)


def _step(
    name: str,
    fn,
    request: dict[str, Any] | None = None,
) -> Step:
    started = time.perf_counter()
    try:
        response, checks = fn()
        ok = all(check.passed for check in checks)
        return Step(
            name=name,
            ok=ok,
            elapsed_seconds=time.perf_counter() - started,
            request=request or {},
            response=_jsonable(response),
            checks=checks,
        )
    except Exception as exc:
        return Step(
            name=name,
            ok=False,
            elapsed_seconds=time.perf_counter() - started,
            request=request or {},
            error=f"{type(exc).__name__}: {exc}",
        )


def _source_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown"}:
        return "markdown"
    if suffix in {".txt", ".text"}:
        return "txt"
    if suffix == ".docx":
        return "docx"
    raise ValueError(f"Unsupported test input suffix: {path.suffix}")


def _content_type(path: Path) -> str:
    return mimetypes.guess_type(path.name)[0] or "application/octet-stream"


def _ensure_markdown_input(path: Path) -> Path:
    if path.exists():
        return path
    GENERATED_INPUT_DIR.mkdir(parents=True, exist_ok=True)
    fallback = GENERATED_INPUT_DIR / "phase2_manual.md"
    fallback.write_text(
        "\n".join(
            [
                "# Product Manual",
                "",
                "Users can upload Markdown, TXT, and DOCX knowledge documents.",
                "After upload, ingestion parses the document and creates searchable chunks.",
                "",
                "## Upload status",
                "",
                "| Status | Meaning |",
                "|---|---|",
                "| processing | The document is being parsed. |",
                "| success | The document is searchable in the knowledge base. |",
                "| failed | The user should inspect the failure and upload again. |",
                "",
                "- Single file size limit is 10 MB",
                "- Batch upload is supported",
                "- Default category is general",
            ]
        ),
        encoding="utf-8",
    )
    return fallback


def _create_docx_input() -> Path:
    GENERATED_INPUT_DIR.mkdir(parents=True, exist_ok=True)
    output = GENERATED_INPUT_DIR / "phase2_manual.docx"

    doc = DocxDocument()
    doc.add_heading("Phase 2 DOCX Manual", level=1)
    doc.add_paragraph(
        "DOCX files should be parsed into titles, paragraphs, lists, tables, and images."
    )
    doc.add_paragraph("Upload the DOCX document.", style="List Bullet")
    doc.add_paragraph("Ingest it into the knowledge base.", style="List Bullet")
    doc.add_paragraph("Search for DOCX content.", style="List Bullet")

    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Area"
    table.cell(0, 1).text = "Capability"
    table.cell(0, 2).text = "Result"
    table.cell(1, 0).text = "DOCX"
    table.cell(1, 1).text = "Table parsing"
    table.cell(1, 2).text = "structured_data.table"
    table.cell(2, 0).text = "DOCX"
    table.cell(2, 1).merge(table.cell(2, 2))
    table.cell(2, 1).text = "Merged cells are expanded"

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
    )
    image_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    try:
        image_file.write(png_bytes)
        image_file.close()
        doc.add_picture(image_file.name)
    finally:
        os.unlink(image_file.name)

    doc.save(output)
    return output


def _health(client: TestClient) -> Step:
    def call():
        resp = client.get("/health")
        body = resp.json()
        return (
            {"status_code": resp.status_code, "body": body},
            [
                Check("HTTP 200", resp.status_code == 200, str(resp.status_code)),
                Check("status ok", body.get("status") == "ok", str(body)),
            ],
        )

    return _step("GET /health", call)


def _upload(client: TestClient, path: Path, title: str, category: str) -> Step:
    request = {
        "file_name": path.name,
        "content_type": _content_type(path),
        "title": title,
        "category": category,
        "size": path.stat().st_size,
    }

    def call():
        with path.open("rb") as handle:
            resp = client.post(
                "/upload",
                files={
                    "file": (
                        path.name,
                        handle,
                        _content_type(path),
                    )
                },
                data={"title": title, "category": category},
            )
        body = resp.json()
        uploaded_path = Path(body.get("source_uri", "").removeprefix("file://"))
        return (
            {"status_code": resp.status_code, "body": body},
            [
                Check("HTTP 200", resp.status_code == 200, str(resp.status_code)),
                Check("source_uri returned", bool(body.get("source_uri")), body.get("source_uri", "")),
                Check("source_hash returned", bool(body.get("source_hash")), body.get("source_hash", "")),
                Check("uploaded file exists", uploaded_path.exists(), str(uploaded_path)),
            ],
        )

    return _step("POST /upload", call, request)


def _ingest(
    client: TestClient,
    upload_body: dict[str, Any],
    source_type: str,
    category: str,
) -> Step:
    payload = {
        "documents": [
            {
                "title": upload_body["title"],
                "source_type": source_type,
                "source_uri": upload_body["source_uri"],
                "category": category,
            }
        ],
        "options": {"max_depth": 1, "extract_assets": True},
    }

    def call():
        resp = client.post("/ingest", json=payload)
        body = resp.json()
        return (
            {"status_code": resp.status_code, "body": body},
            [
                Check("HTTP 202", resp.status_code == 202, str(resp.status_code)),
                Check("accepted", body.get("status") == "accepted", str(body)),
                Check("job_id returned", bool(body.get("job_id")), str(body.get("job_id"))),
            ],
        )

    return _step("POST /ingest", call, payload)


def _poll(
    client: TestClient,
    job_id: str,
    timeout_seconds: float,
    require_chunks: bool,
    min_assets: int,
) -> Step:
    request = {
        "job_id": job_id,
        "timeout_seconds": timeout_seconds,
        "require_chunks": require_chunks,
        "min_assets": min_assets,
    }

    def call():
        deadline = time.time() + timeout_seconds
        polls: list[dict[str, Any]] = []
        final: dict[str, Any] | None = None
        while time.time() < deadline:
            resp = client.get(f"/ingest/{job_id}")
            body = resp.json()
            item = {"status_code": resp.status_code, **body}
            polls.append(item)
            final = item
            if body.get("status") in {"completed", "failed"}:
                break
            time.sleep(0.5)

        final = final or {}
        chunk_count = int(final.get("chunk_count") or 0)
        asset_count = int(final.get("asset_count") or 0)
        return (
            {"polls": polls, "final": final},
            [
                Check("status endpoint HTTP 200", final.get("status_code") == 200, str(final.get("status_code"))),
                Check("job completed", final.get("status") == "completed", str(final.get("status"))),
                Check("job has no error", not final.get("error"), str(final.get("error"))),
                Check("chunk_count > 0", (chunk_count > 0) if require_chunks else True, str(chunk_count)),
                Check("asset_count >= minimum", asset_count >= min_assets, str(asset_count)),
            ],
        )

    return _step("GET /ingest/{job_id}", call, request)


def _search(
    client: TestClient,
    query: str,
    category: str,
    require_results: bool,
) -> Step:
    payload = {"query": query, "top_k": 5, "filters": {"category": category}}

    def call():
        resp = client.post("/search", json=payload)
        body = resp.json()
        results = body.get("results") or []
        return (
            {"status_code": resp.status_code, "body": body},
            [
                Check("HTTP 200", resp.status_code == 200, str(resp.status_code)),
                Check("search_id returned", bool(body.get("search_id")), body.get("search_id", "")),
                Check("results returned", bool(results) if require_results else True, f"results={len(results)}"),
                Check(
                    "score details present",
                    all("score_components" in item for item in results) if results else not require_results,
                    str(results[0].get("score_components", {})) if results else "{}",
                ),
                Check(
                    "source refs present",
                    all(item.get("source_refs") for item in results) if results else not require_results,
                    str(results[0].get("source_refs", [])) if results else "[]",
                ),
            ],
        )

    return _step("POST /search", call, payload)


def _poll_chunk_count(step: Step) -> int:
    if not isinstance(step.response, dict):
        return 0
    final = step.response.get("final") or {}
    return int(final.get("chunk_count") or 0)


def _run_case(
    client: TestClient,
    name: str,
    input_path: Path,
    query: str,
    category: str,
    timeout_seconds: float,
    require_chunks: bool,
) -> CaseResult:
    source_type = _source_type(input_path)
    title = f"{name}-{datetime.now().strftime('%H%M%S')}"
    steps: list[Step] = []

    steps.append(_health(client))
    upload_step = _upload(client, input_path, title, category)
    steps.append(upload_step)

    upload_body = {}
    if isinstance(upload_step.response, dict):
        upload_body = upload_step.response.get("body") or {}

    if upload_step.ok:
        ingest_step = _ingest(client, upload_body, source_type, category)
    else:
        ingest_step = Step("POST /ingest", False, error="Skipped because upload failed")
    steps.append(ingest_step)

    ingest_body = {}
    if isinstance(ingest_step.response, dict):
        ingest_body = ingest_step.response.get("body") or {}
    job_id = ingest_body.get("job_id")

    poll_step: Step
    if ingest_step.ok and job_id:
        min_assets = 1 if source_type == "docx" else 0
        poll_step = _poll(client, str(job_id), timeout_seconds, require_chunks, min_assets)
    else:
        poll_step = Step("GET /ingest/{job_id}", False, error="Skipped because ingest failed")
    steps.append(poll_step)

    if not poll_step.ok or _poll_chunk_count(poll_step) <= 0:
        steps.append(
            Step(
                "POST /search",
                False if require_chunks else True,
                request={"query": query, "category": category},
                error="Skipped because ingestion did not produce chunks",
                checks=[
                    Check(
                        "search skipped after empty ingestion",
                        not require_chunks,
                        f"chunk_count={_poll_chunk_count(poll_step)}",
                    )
                ],
            )
        )
    else:
        steps.append(_search(client, query, category, require_results=require_chunks))
    return CaseResult(name=name, input_path=str(input_path), source_type=source_type, steps=steps)


def _status_text(ok: bool) -> str:
    return "PASS" if ok else "FAIL"


def _render_markdown(
    *,
    started_at: str,
    finished_at: str,
    backend: str,
    api_key_configured: bool,
    cases: list[CaseResult],
) -> str:
    overall_ok = all(case.ok for case in cases)
    lines = [
        "# Phase 2 End-to-End Test Report",
        "",
        f"- Started: {started_at}",
        f"- Finished: {finished_at}",
        f"- Overall: **{_status_text(overall_ok)}**",
        f"- Backend: `{backend}`",
        f"- VOLCENGINE_API_KEY configured: `{api_key_configured}`",
        f"- Cases: {sum(1 for case in cases if case.ok)}/{len(cases)} passed",
        "",
        "## Summary",
        "",
        "| Case | Source Type | Result |",
        "|---|---|---|",
    ]
    for case in cases:
        lines.append(f"| {case.name} | `{case.source_type}` | {_status_text(case.ok)} |")

    for case in cases:
        lines.extend(["", f"## Case: {case.name}", "", f"- Input: `{case.input_path}`", ""])
        for idx, step in enumerate(case.steps, 1):
            lines.extend(
                [
                    f"### {idx}. {step.name}",
                    "",
                    f"- Result: **{_status_text(step.ok)}**",
                    f"- Elapsed: {step.elapsed_seconds:.2f}s",
                ]
            )
            if step.error:
                lines.append(f"- Error: `{step.error}`")
            lines.extend(["", "#### Checks", "", "| Check | Result | Detail |", "|---|---|---|"])
            if step.checks:
                for check in step.checks:
                    detail = str(check.detail).replace("\n", " ")[:300]
                    lines.append(f"| {check.name} | {_status_text(check.passed)} | `{detail}` |")
            else:
                lines.append("| No checks | FAIL | `No checks recorded` |")
            lines.extend(["", "<details><summary>Request</summary>", "", "```json"])
            lines.append(json.dumps(step.request, ensure_ascii=False, indent=2))
            lines.extend(["```", "", "</details>", "", "<details><summary>Response</summary>", "", "```json"])
            lines.append(json.dumps(step.response, ensure_ascii=False, indent=2))
            lines.extend(["```", "", "</details>"])

    return "\n".join(lines) + "\n"


def _write_reports(cases: list[CaseResult], started_at: str, backend: str, api_key_configured: bool) -> tuple[Path, Path]:
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    finished_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    markdown_path = RESULTS_DIR / f"phase2_e2e_{stamp}.md"
    json_path = RESULTS_DIR / f"phase2_e2e_{stamp}.json"

    markdown_path.write_text(
        _render_markdown(
            started_at=started_at,
            finished_at=finished_at,
            backend=backend,
            api_key_configured=api_key_configured,
            cases=cases,
        ),
        encoding="utf-8",
    )
    json_path.write_text(
        json.dumps(
            {
                "started_at": started_at,
                "finished_at": finished_at,
                "overall_ok": all(case.ok for case in cases),
                "backend": backend,
                "volcengine_api_key_configured": api_key_configured,
                "cases": [
                    {
                        "name": case.name,
                        "input_path": case.input_path,
                        "source_type": case.source_type,
                        "ok": case.ok,
                        "steps": [
                            {
                                "name": step.name,
                                "ok": step.ok,
                                "elapsed_seconds": step.elapsed_seconds,
                                "request": step.request,
                                "response": step.response,
                                "checks": [check.__dict__ for check in step.checks],
                                "error": step.error,
                            }
                            for step in case.steps
                        ],
                    }
                    for case in cases
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return markdown_path, json_path


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--markdown-input", type=Path, default=DEFAULT_MARKDOWN_INPUT)
    parser.add_argument("--timeout", type=float, default=60.0)
    parser.add_argument("--category", default=DEFAULT_CATEGORY)
    parser.add_argument(
        "--allow-zero-chunks",
        action="store_true",
        help="Do not fail chunk/search checks when VOLCENGINE_API_KEY is unavailable.",
    )
    args = parser.parse_args()

    os.chdir(PACKAGE_ROOT)

    from app.core.config import settings
    from app.main import app

    started_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    markdown_input = _ensure_markdown_input(args.markdown_input.resolve())
    docx_input = _create_docx_input()
    require_chunks = not args.allow_zero_chunks

    client = TestClient(app)
    cases = [
        _run_case(
            client,
            "markdown-chain",
            markdown_input,
            "How can a user confirm an uploaded document was parsed successfully?",
            args.category,
            args.timeout,
            require_chunks,
        ),
        _run_case(
            client,
            "docx-chain",
            docx_input,
            "What DOCX capabilities are parsed into the knowledge base?",
            args.category,
            args.timeout,
            require_chunks,
        ),
    ]

    markdown_path, json_path = _write_reports(
        cases,
        started_at,
        settings.backend,
        bool(settings.api_key),
    )
    print(f"Markdown report: {markdown_path}")
    print(f"JSON report: {json_path}")
    print(f"Overall: {_status_text(all(case.ok for case in cases))}")
    return 0 if all(case.ok for case in cases) else 1


if __name__ == "__main__":
    raise SystemExit(main())
