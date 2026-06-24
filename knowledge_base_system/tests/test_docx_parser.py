import io
import base64
import os
import tempfile

import pytest
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.models import AssetType, Document, ElementType
from parsers.docx_parser import DocxParser


def _make_test_docx() -> bytes:
    """Create a minimal DOCX with heading, paragraph, and table."""
    doc = DocxDocument()

    doc.add_heading("测试文档", level=1)
    doc.add_paragraph("这是一个测试段落，用于验证 DOCX 解析器。")

    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    # Header row
    table.cell(0, 0).text = "状态"
    table.cell(0, 1).text = "说明"
    # Data rows
    table.cell(1, 0).text = "处理中"
    table.cell(1, 1).text = "系统正在解析文档"
    table.cell(2, 0).text = "成功"
    table.cell(2, 1).text = "文档已进入知识库"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxParser:
    def setup_method(self):
        self.parser = DocxParser()

    def test_supported_types(self):
        assert self.parser.supports("docx")
        assert self.parser.SUPPORTED_TYPES == {"docx"}
        assert not self.parser.supports("pdf")

    def test_parse_headings_and_paragraphs(self):
        content = _make_test_docx()
        doc = Document(
            title="Test",
            source_type="docx",
            source_uri="memory://test",
            metadata={"raw_content": content},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        titles = [e for e in result.elements if e.element_type == ElementType.title]
        paragraphs = [e for e in result.elements if e.element_type == ElementType.paragraph]

        assert len(titles) >= 1
        assert titles[0].text == "测试文档"
        assert titles[0].metadata.get("heading_level") == 1
        assert len(paragraphs) >= 1

    def test_parse_table(self):
        content = _make_test_docx()
        doc = Document(
            title="Test",
            source_type="docx",
            source_uri="memory://test",
            metadata={"raw_content": content},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        tables = [e for e in result.elements if e.element_type == ElementType.table]
        assert len(tables) >= 1

        table = tables[0]
        assert table.structured_data is not None
        table_data = table.structured_data["table"]
        # headers 和 cells 现在是 {"text": ..., "asset_data": [...]} 格式
        assert [h["text"] for h in table_data["headers"]] == ["状态", "说明"]
        assert len(table_data["rows"]) == 2

    def test_document_hash_set(self):
        content = _make_test_docx()
        doc = Document(
            title="Test",
            source_type="docx",
            source_uri="memory://test",
            metadata={"raw_content": content},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])
        assert result.doc.source_hash
        assert result.doc.source_hash.startswith("sha256:")

    def test_all_elements_have_sequence_order(self):
        content = _make_test_docx()
        doc = Document(
            title="Test",
            source_type="docx",
            source_uri="memory://test",
            metadata={"raw_content": content},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        orders = [e.sequence_order for e in result.elements]
        assert all(o > 0 for o in orders)
        assert len(set(orders)) == len(orders)

    def test_returns_parse_result(self):
        content = _make_test_docx()
        doc = Document(
            title="Test",
            source_type="docx",
            source_uri="memory://test",
            metadata={"raw_content": content},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])
        assert result.doc == doc
        assert len(result.elements) > 0
        for el in result.elements:
            assert el.doc_id == doc.doc_id

    def test_parse_list_items(self):
        docx = DocxDocument()
        docx.add_paragraph("First item", style="List Bullet")
        docx.add_paragraph("Second item", style="List Bullet")
        buf = io.BytesIO()
        docx.save(buf)

        doc = Document(
            title="List",
            source_type="docx",
            source_uri="memory://list",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        lists = [e for e in result.elements if e.element_type == ElementType.list]
        items = [
            e for e in result.elements
            if e.element_type == ElementType.paragraph and e.parent_element_id
        ]
        assert len(lists) == 1
        assert [item.text for item in items] == ["First item", "Second item"]
        assert all(item.parent_element_id == lists[0].element_id for item in items)

    def test_parse_merged_table_cells_are_expanded(self):
        docx = DocxDocument()
        table = docx.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(0, 2).text = "C"
        table.cell(1, 0).text = "Merged"
        table.cell(1, 0).merge(table.cell(1, 1))
        table.cell(1, 2).text = "Tail"
        buf = io.BytesIO()
        docx.save(buf)

        doc = Document(
            title="Merged",
            source_type="docx",
            source_uri="memory://merged",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])
        table_el = next(e for e in result.elements if e.element_type == ElementType.table)
        row = table_el.structured_data["table"]["rows"][0]["cells"]
        assert [cell["text"] for cell in row] == ["Merged", "Merged", "Tail"]

    def test_extract_embedded_image_from_raw_content(self):
        png_bytes = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
        docx = DocxDocument()
        image_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        try:
            image_file.write(png_bytes)
            image_file.close()
            docx.add_picture(image_file.name)
            buf = io.BytesIO()
            docx.save(buf)
        finally:
            os.unlink(image_file.name)

        doc = Document(
            title="Image",
            source_type="docx",
            source_uri="memory://image",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        assert len(result.assets) == 1
        assert result.assets[0].content_hash.startswith("sha256:")
        # 图片属于附属资源，不创建独立元素
        images = [e for e in result.elements if "[图片" in e.text]
        assert len(images) >= 1  # 图片信息作为 paragraph 存在
        # 图片作为段落资源关联
        paragraphs = [e for e in result.elements if e.element_type == ElementType.paragraph]
        assert len(paragraphs) >= 1
        # 至少有一个段落关联了图片 asset
        para_with_image = [p for p in paragraphs if p.asset_data]
        assert len(para_with_image) >= 1
        assert any(ad.url == result.assets[0].original_uri for ad in para_with_image[0].asset_data)

    def test_unsupported_embedded_object_degrades_to_unknown(self):
        docx = DocxDocument()
        paragraph = docx.add_paragraph()
        paragraph._p.append(OxmlElement("w:object"))
        buf = io.BytesIO()
        docx.save(buf)

        doc = Document(
            title="Unknown",
            source_type="docx",
            source_uri="memory://unknown",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        unknowns = [e for e in result.elements if e.element_type == ElementType.unknown]
        assert len(unknowns) == 1
        assert "不支持的嵌入" in unknowns[0].text

    # ── 新增测试：非英文标题 ────────────────────────────────────

    def test_heading_chinese(self):
        """中文"标题 1"样式正确识别。"""
        docx = DocxDocument()
        p = docx.add_paragraph("中文标题内容")
        # 模拟中文 Word 的样式名
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._p.insert(0, pPr)
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), "标题 1")
        pPr.append(pStyle)
        buf = io.BytesIO()
        docx.save(buf)

        doc = Document(
            title="CN", source_type="docx", source_uri="memory://cn",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])
        titles = [e for e in result.elements if e.element_type == ElementType.title]
        assert len(titles) == 1
        assert titles[0].metadata.get("heading_level") == 1
        assert titles[0].text == "中文标题内容"

    def test_heading_french(self):
        """法文"Titre 1"样式正确识别。"""
        docx = DocxDocument()
        p = docx.add_paragraph("Contenu du titre")
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._p.insert(0, pPr)
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), "Titre 1")
        pPr.append(pStyle)
        buf = io.BytesIO()
        docx.save(buf)

        doc = Document(
            title="FR", source_type="docx", source_uri="memory://fr",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])
        titles = [e for e in result.elements if e.element_type == ElementType.title]
        assert len(titles) == 1
        assert titles[0].metadata.get("heading_level") == 1

    # ── 新增测试：段落内联图片 ──────────────────────────────────

    def test_paragraph_with_image(self):
        """段落中内联图片关联到 asset_data，文本含 [图片: xxx]。"""
        png_bytes = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
        docx = DocxDocument()
        image_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        try:
            image_file.write(png_bytes)
            image_file.close()
            docx.add_picture(image_file.name)
            buf = io.BytesIO()
            docx.save(buf)
        finally:
            os.unlink(image_file.name)

        doc = Document(
            title="Img", source_type="docx", source_uri="memory://img",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        paragraphs = [e for e in result.elements if e.element_type == ElementType.paragraph]
        assert len(paragraphs) >= 1
        # 段落文本包含 [图片: ...] 占位符
        img_para = paragraphs[0]
        assert "[图片:" in img_para.text
        # 图片关联到段落的 asset_data
        assert len(img_para.asset_data) >= 1
        assert len(result.assets) >= 1

    def test_image_only_paragraph(self):
        """纯图片段落（无文本）仍创建元素。"""
        png_bytes = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
        docx = DocxDocument()
        image_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        try:
            image_file.write(png_bytes)
            image_file.close()
            docx.add_picture(image_file.name)
            buf = io.BytesIO()
            docx.save(buf)
        finally:
            os.unlink(image_file.name)

        doc = Document(
            title="ImgOnly", source_type="docx", source_uri="memory://imgonly",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        # 纯图片段落应创建段落元素
        paragraphs = [e for e in result.elements if e.element_type == ElementType.paragraph]
        assert len(paragraphs) >= 1

    # ── 新增测试：段落超链接 ────────────────────────────────────

    def test_paragraph_with_hyperlink(self):
        """段落中普通网页超链接 URL 写入 metadata.link_urls，文本保留显示文字。"""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        docx = DocxDocument()
        p = docx.add_paragraph()
        run = p.add_run("访问 ")
        rel = docx.part.relate_to("https://www.example.com", RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rel)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "示例网站"
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        hyperlink.append(r)
        p._p.append(hyperlink)
        run2 = p.add_run(" 了解更多")
        buf = io.BytesIO()
        docx.save(buf)

        doc = Document(
            title="Link", source_type="docx", source_uri="memory://link",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        paragraphs = [e for e in result.elements if e.element_type == ElementType.paragraph]
        assert len(paragraphs) >= 1
        para = paragraphs[0]
        # 文本保留显示文字
        assert "示例网站" in para.text
        assert "了解更多" in para.text
        # URL 写入 metadata.link_urls
        assert "link_urls" in para.metadata
        assert "https://www.example.com" in para.metadata["link_urls"]

    def test_paragraph_with_attachment_link(self):
        """文件/附件链接创建 Asset，段落文本只保留显示文字。"""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        docx = DocxDocument()
        p = docx.add_paragraph()
        rel = docx.part.relate_to("https://example.com/report.pdf", RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rel)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "下载报告"
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        hyperlink.append(r)
        p._p.append(hyperlink)
        buf = io.BytesIO()
        docx.save(buf)

        doc = Document(
            title="Attach", source_type="docx", source_uri="memory://attach",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        # 应有 attachment 类型 Asset
        attach_assets = [a for a in result.assets if a.asset_type == AssetType.document_link]
        assert len(attach_assets) >= 1
        # 段落关联了 attachment
        paragraphs = [e for e in result.elements if e.element_type == ElementType.paragraph]
        assert len(paragraphs) >= 1
        assert any(ad.url == attach_assets[0].original_uri for ad in paragraphs[0].asset_data)
        # 文本只保留显示文字
        assert "下载报告" in paragraphs[0].text

    def test_paragraph_with_video_link(self):
        """视频链接创建 Asset 并关联到段落。"""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        docx = DocxDocument()
        p = docx.add_paragraph()
        rel = docx.part.relate_to("https://www.youtube.com/watch?v=abc123", RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rel)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "观看视频"
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        hyperlink.append(r)
        p._p.append(hyperlink)
        buf = io.BytesIO()
        docx.save(buf)

        doc = Document(
            title="Video", source_type="docx", source_uri="memory://video",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        video_assets = [a for a in result.assets if a.asset_type == AssetType.video_link]
        assert len(video_assets) >= 1
        assert video_assets[0].original_uri == "https://www.youtube.com/watch?v=abc123"

    # ── 新增测试：表格单元格资源 ────────────────────────────────

    def test_table_cell_with_image(self):
        """表格单元格图片关联到 structured_data 的 cell asset_data。"""
        png_bytes = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
        image_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        image_file.write(png_bytes)
        image_file.close()

        docx = DocxDocument()
        table = docx.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "文本格"
        # 在 (0,1) 插入图片
        p = table.cell(0, 1).paragraphs[0]
        p.clear()
        run = p.add_run()
        run.add_picture(image_file.name)
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"
        buf = io.BytesIO()
        docx.save(buf)
        os.unlink(image_file.name)

        doc = Document(
            title="TCImg", source_type="docx", source_uri="memory://tcimg",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        tables = [e for e in result.elements if e.element_type == ElementType.table]
        assert len(tables) == 1
        table_data = tables[0].structured_data["table"]
        # 图片在 header 行 (0,1)，不在 data 行
        assert table_data["rows"][0]["cells"][0]["asset_data"] == []
        assert len(table_data["headers"][1]["asset_data"]) >= 1
        # 表格级 asset_data 汇总
        assert len(tables[0].asset_data) >= 1

    def test_table_cell_with_hyperlink(self):
        """表格单元格超链接正确处理。"""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        docx = DocxDocument()
        table = docx.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "表头A"
        table.cell(0, 1).text = "表头B"
        # 在 (1,1) 添加附件超链接（数据行）
        rel = docx.part.relate_to("https://example.com/doc.pdf", RT.HYPERLINK, is_external=True)
        p = table.cell(1, 1).paragraphs[0]
        p.clear()
        run = p.add_run("下载 ")
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rel)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "文档"
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        hyperlink.append(r)
        p._p.append(hyperlink)
        buf = io.BytesIO()
        docx.save(buf)

        doc = Document(
            title="TCLink", source_type="docx", source_uri="memory://tclink",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        tables = [e for e in result.elements if e.element_type == ElementType.table]
        assert len(tables) == 1
        table_data = tables[0].structured_data["table"]
        # (1,1) 的 asset_data 应包含 attachment
        cell_asset_data = table_data["rows"][0]["cells"][1]["asset_data"]
        assert len(cell_asset_data) >= 1
        # 文本包含显示文字
        assert "文档" in table_data["rows"][0]["cells"][1]["text"]
        # 应有 attachment Asset
        attach_assets = [a for a in result.assets if a.asset_type == AssetType.document_link]
        assert len(attach_assets) >= 1

    # ── 新增测试：合并单元格 + 资源 ─────────────────────────────

    def test_merged_cell_with_asset(self):
        """gridSpan 合并单元格正确传递 asset_data。"""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        docx = DocxDocument()
        table = docx.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "H1"
        table.cell(0, 1).text = "H2"
        table.cell(1, 0).text = "A"
        # (1,1) 添加图片链接（数据行）
        rel = docx.part.relate_to("https://example.com/photo.png", RT.HYPERLINK, is_external=True)
        p = table.cell(1, 1).paragraphs[0]
        p.clear()
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rel)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "图片"
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        hyperlink.append(r)
        p._p.append(hyperlink)
        # 第三行合并
        table.cell(2, 0).text = "Merged"
        table.cell(2, 0).merge(table.cell(2, 1))
        buf = io.BytesIO()
        docx.save(buf)

        doc = Document(
            title="MergeAsset", source_type="docx", source_uri="memory://mergeasset",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        tables = [e for e in result.elements if e.element_type == ElementType.table]
        assert len(tables) == 1
        table_data = tables[0].structured_data["table"]
        # 第一行数据 (1,1) 应有图片 asset
        assert len(table_data["rows"][0]["cells"][1]["asset_data"]) >= 1

    # ── 新增测试：资源回填 ──────────────────────────────────────

    def test_asset_source_element_backfill(self):
        """_link_assets_to_elements 正确回填 element_id。"""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        docx = DocxDocument()
        p = docx.add_paragraph()
        rel = docx.part.relate_to("https://example.com/doc.pdf", RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rel)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "下载"
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        hyperlink.append(r)
        p._p.append(hyperlink)
        buf = io.BytesIO()
        docx.save(buf)

        doc = Document(
            title="Backfill", source_type="docx", source_uri="memory://backfill",
            metadata={"raw_content": buf.getvalue()},
        )
        result = self.parser.parse(doc, doc.metadata["raw_content"])

        # 所有有 element_id 的 Asset 应指向存在的 element_id
        element_ids = {el.element_id for el in result.elements}
        for asset in result.assets:
            if asset.element_id:
                assert asset.element_id in element_ids
